from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:  # pragma: no cover
    Image = None


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "AI Violence Detection System Documentation.docx"
DIAGRAM = ROOT / "documentation_architecture.png"


BLUE = RGBColor(23, 54, 93)
GOLD = RGBColor(197, 146, 42)
LIGHT_BLUE = "D9EAF7"
LIGHT_GOLD = "F8E9C0"
LIGHT_GRAY = "F2F2F2"
TEXT = RGBColor(33, 33, 33)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_field(paragraph, instr):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_toc(doc, title, field_code):
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    add_field(p, field_code)
    p.add_run("Right-click and update this field in Microsoft Word if page numbers need refreshing.")
    p.runs[-1].italic = True
    p.runs[-1].font.size = Pt(10)
    doc.add_page_break()


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = "Caption"
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, color=BLUE)
        set_cell_shading(hdr[i], LIGHT_BLUE)
        if widths:
            set_cell_width(hdr[i], widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                set_cell_width(cells[i], widths[i])
    doc.add_paragraph()
    return table


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.text = "AI Violence Detection System"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Page ")
    add_field(p, "PAGE")
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(100, 100, 100)


def style_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.first_line_indent = Inches(0.25)

    for name, size, color in [
        ("Title", 22, BLUE),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 12, BLUE),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.first_line_indent = Inches(0)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)

    cap = styles["Caption"]
    cap.font.name = "Arial"
    cap.font.size = Pt(10)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor(80, 80, 80)
    cap.paragraph_format.first_line_indent = Inches(0)


def draw_architecture():
    if Image is None:
        return None
    width, height = 1700, 760
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    try:
        title_font = ImageFont.truetype("arialbd.ttf", 34)
        box_font = ImageFont.truetype("arialbd.ttf", 24)
        small_font = ImageFont.truetype("arial.ttf", 18)
    except Exception:
        title_font = box_font = small_font = None

    draw.rectangle([0, 0, width, 90], fill=(23, 54, 93))
    draw.text((40, 26), "High-Level Architecture of the AI Violence Detection System", fill="white", font=title_font)

    boxes = [
        (70, 185, 330, 335, "CCTV / Camera\nSources", "Live video frames from monitored areas"),
        (430, 130, 760, 280, "Weapon AI Service", "OpenCV + YOLO model\nDetects visible weapons"),
        (430, 390, 760, 540, "Violence AI Service", "OpenCV + Keras model\nClassifies violent activity"),
        (880, 250, 1180, 420, "Laravel Backend\n+ MySQL", "Stores cameras, incidents,\nalerts, snapshots, clips"),
        (1320, 250, 1625, 420, "React Operator\nDashboard", "Live stream, incidents,\nevidence, reports"),
    ]
    for x1, y1, x2, y2, title, desc in boxes:
        draw.rounded_rectangle([x1, y1, x2, y2], radius=22, fill=(217, 234, 247), outline=(23, 54, 93), width=4)
        title_lines = title.split("\n")
        y = y1 + 24
        for line in title_lines:
            tw = draw.textlength(line, font=box_font)
            draw.text((x1 + (x2 - x1 - tw) / 2, y), line, fill=(23, 54, 93), font=box_font)
            y += 30
        y += 12
        for line in desc.split("\n"):
            tw = draw.textlength(line, font=small_font)
            draw.text((x1 + (x2 - x1 - tw) / 2, y), line, fill=(40, 40, 40), font=small_font)
            y += 24

    def arrow(start, end, label=None, label_pos=None):
        draw.line([start, end], fill=(197, 146, 42), width=6)
        ex, ey = end
        sx, sy = start
        if ex >= sx:
            pts = [(ex, ey), (ex - 22, ey - 12), (ex - 22, ey + 12)]
        else:
            pts = [(ex, ey), (ex + 22, ey - 12), (ex + 22, ey + 12)]
        draw.polygon(pts, fill=(197, 146, 42))
        if label and label_pos:
            draw.text(label_pos, label, fill=(90, 70, 20), font=small_font)

    arrow((330, 255), (430, 205), "Frames", (345, 205))
    arrow((330, 255), (430, 465), "Frames", (345, 455))
    arrow((760, 205), (880, 300), "POST incidents", (770, 235))
    arrow((760, 465), (880, 370), "POST incidents", (770, 455))
    arrow((1180, 335), (1320, 335), "Review and respond", (1195, 302))

    draw.rectangle([70, 630, 1625, 700], fill=(248, 233, 192), outline=(197, 146, 42), width=2)
    draw.text(
        (95, 652),
        "Evidence flow: detected event -> confidence/alert level -> snapshot and video clip -> stored incident record -> operator inspection.",
        fill=(60, 45, 20),
        font=small_font,
    )
    img.save(DIAGRAM)
    return DIAGRAM


def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("Ain Shams University\nFaculty of Computer and Information Sciences\nComputer Systems Department")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(14)
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI Violence Detection System")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(28)
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Graduation Project Documentation")
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.color.rgb = GOLD

    doc.add_paragraph()
    tbl = add_table(
        doc,
        ["Item", "Details"],
        [
            ["Project Area", "Artificial Intelligence, Computer Vision, Public Safety, Web Monitoring Systems"],
            ["Project Scope", "Real-time violence and weapon detection from camera streams with incident recording and dashboard review"],
            ["Academic Year", "2025/2026"],
        ],
        widths=[2500, 6500],
    )
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    r = p.add_run("Prepared by\nProject Team Members")
    r.font.size = Pt(12)
    r.font.name = "Arial"
    r.bold = True
    doc.add_page_break()

    doc.add_heading("Internal Cover Page", level=1)
    rows = [
        ["University", "Ain Shams University"],
        ["Faculty", "Faculty of Computer and Information Sciences"],
        ["Department", "Computer Systems Department"],
        ["Project Title", "AI Violence Detection System"],
        ["Submitted By", "Project Team Members"],
        ["Supervisor", "Project Supervisor"],
        ["Year", "2025/2026"],
    ]
    add_table(doc, ["Field", "Information"], rows, widths=[2500, 6500])
    doc.add_page_break()


def front_matter(doc):
    doc.add_heading("Acknowledgement", level=1)
    doc.add_paragraph(
        "We would like to express our sincere gratitude to our supervisor, faculty members, teaching assistants, and everyone who supported this graduation project. "
        "Their guidance helped us connect artificial intelligence concepts with a practical public-safety application. We also thank our families and colleagues for their encouragement throughout the design, implementation, and testing phases."
    )
    doc.add_page_break()

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "Violence and weapon-related incidents can escalate within seconds, especially in crowded public places, transportation areas, campuses, commercial zones, and sensitive facilities. "
        "Traditional surveillance depends mainly on continuous human observation, which becomes difficult when many camera feeds must be monitored at the same time. The AI Violence Detection System addresses this problem by using computer vision to analyze camera streams, detect violent behavior or weapons, record evidence, and notify operators through a centralized web dashboard."
    )
    doc.add_paragraph(
        "The project is motivated by the need to support safer communities locally across Egypt and globally wherever camera-based monitoring is used. The system does not replace law-enforcement judgment; instead, it provides faster awareness, organized evidence, and a structured incident workflow that can help prevent criminal actions from developing unnoticed. The proposed solution combines Python/OpenCV AI services, a YOLO-based weapon detection model, a Keras-based violence classification model, a Laravel backend, a MySQL database, and a React dashboard."
    )
    doc.add_paragraph(
        "The solution demonstrates how local camera sources can be connected to AI analysis services that generate incident records with confidence scores, alert levels, snapshots, and video clips. Operators can then inspect incidents, review camera information, and use stored evidence for response and reporting. The system is evaluated through functional testing of stream access, incident submission, dashboard display, evidence storage, and alert classification."
    )
    doc.add_page_break()

    add_toc(doc, "Table of Contents", 'TOC \\o "1-3" \\h \\z \\u')
    add_toc(doc, "List of Figures", 'TOC \\h \\z \\c "Figure"')
    add_toc(doc, "List of Tables", 'TOC \\h \\z \\c "Table"')

    doc.add_heading("List of Abbreviations", level=1)
    add_table(
        doc,
        ["Abbreviation", "Meaning"],
        [
            ["AI", "Artificial Intelligence"],
            ["API", "Application Programming Interface"],
            ["CCTV", "Closed-Circuit Television"],
            ["CNN", "Convolutional Neural Network"],
            ["CV", "Computer Vision"],
            ["HTTP", "Hypertext Transfer Protocol"],
            ["JSON", "JavaScript Object Notation"],
            ["MVC", "Model-View-Controller"],
            ["OpenCV", "Open Source Computer Vision Library"],
            ["YOLO", "You Only Look Once object detection model family"],
        ],
        widths=[2200, 7000],
    )
    doc.add_page_break()


def chapter1(doc):
    doc.add_heading("Chapter 1: Introduction", level=1)
    doc.add_heading("1.1 Project Area", level=2)
    doc.add_paragraph(
        "The project belongs to the areas of artificial intelligence, computer vision, public-safety systems, and web-based monitoring dashboards. Computer vision allows software to interpret visual information from images and video frames. In the context of surveillance, it can help detect patterns that may indicate risk, such as aggressive physical actions or visible weapons."
    )
    doc.add_paragraph(
        "Many organizations already use cameras for security, but the value of these cameras depends on the ability to observe and interpret feeds at the right time. When many streams are active, human operators may miss short or unexpected events. AI-assisted monitoring can reduce this delay by analyzing frames continuously and raising an alert when suspicious activity is detected."
    )

    doc.add_heading("1.2 Project Motivation", level=2)
    doc.add_paragraph(
        "The main motivation of this project is to contribute to violence prevention and crime reduction through early detection. Locally, Egypt has many crowded environments such as universities, metro stations, streets, malls, hospitals, and public-service buildings where rapid awareness of violent behavior can help protect citizens and support security teams. A system that highlights risky events quickly can help operators respond before an incident becomes more dangerous."
    )
    doc.add_paragraph(
        "Globally, the same need exists in many countries and communities. Violence, weapon threats, and criminal actions are not limited to one region; they are public-safety challenges that require faster information, better monitoring, and reliable evidence. This project is therefore designed as a locally useful and globally relevant system: it can be adapted to different camera sources, locations, and operational policies."
    )
    doc.add_paragraph(
        "The project is also motivated by the practical limitations of manual surveillance. A human operator can become tired, distracted, or overloaded when watching several feeds for long periods. AI can assist by continuously scanning frames, detecting high-risk patterns, and organizing the resulting evidence. This improves situational awareness while keeping final decisions in the hands of authorized people."
    )

    doc.add_heading("1.3 Problem Definition", level=2)
    doc.add_paragraph(
        "Security teams often depend on live CCTV monitoring to discover violent incidents or weapon threats. However, manual monitoring is slow, difficult to scale, and vulnerable to missed events. A violent action may happen in a few seconds, and if the operator does not notice it immediately, the opportunity for early intervention may be lost."
    )
    doc.add_paragraph(
        "The problem addressed by this project is the absence of an integrated system that can analyze camera streams in real time, detect violence or weapons, classify the seriousness of the event, save related evidence, and present the incident clearly to an operator. Without this integration, camera feeds, AI results, evidence files, and incident records may remain disconnected, making response and investigation less efficient."
    )
    doc.add_paragraph(
        "Therefore, the project focuses on building a practical AI-supported monitoring platform that receives live frames, applies detection models, sends incident data to a backend API, stores evidence, and displays alerts through a dashboard. The goal is to reduce detection delay and improve the organization of information needed to prevent violent and criminal actions."
    )

    doc.add_heading("1.4 Project Objectives", level=2)
    add_numbered(
        doc,
        [
            "Develop an AI-assisted system capable of detecting violence and weapon-related threats from camera streams.",
            "Support local public-safety needs in Egypt while keeping the design adaptable for global deployment contexts.",
            "Provide a web dashboard where operators can monitor incidents, cameras, alert levels, evidence, and recent trends.",
            "Store incident data in a structured database with related confidence scores, timestamps, camera identifiers, districts, snapshots, and video clips.",
            "Integrate separate AI services with a Laravel backend using clear API communication.",
            "Reduce the dependency on continuous manual observation by automatically highlighting suspicious events.",
            "Create a documentation structure that explains the problem, motivation, design, implementation, testing, conclusion, and future work clearly.",
        ],
    )

    doc.add_heading("1.5 System Description and Main Phases", level=2)
    doc.add_paragraph(
        "The AI Violence Detection System is composed of camera sources, Python AI services, a backend API, persistent storage, and an operator dashboard. Each part has a specific role in the detection and response workflow."
    )
    if DIAGRAM.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(DIAGRAM), width=Inches(6.4))
        add_caption(doc, "Figure 1.1: High-level architecture of the AI Violence Detection System")

    doc.add_heading("1.5.1 Camera Input Phase", level=3)
    doc.add_paragraph(
        "The system starts with camera or CCTV sources. In the current implementation, the AI services can use local camera sources during development. These sources provide frames that are read through OpenCV and prepared for analysis."
    )
    doc.add_heading("1.5.2 AI Analysis Phase", level=3)
    doc.add_paragraph(
        "Two AI services analyze the incoming frames. The weapon detection service uses a YOLO model to identify visible weapons, while the violence detection service uses a Keras model to classify violent activity. Each service exposes stream and control endpoints and can generate an incident when the detection threshold is reached."
    )
    doc.add_heading("1.5.3 Incident Recording Phase", level=3)
    doc.add_paragraph(
        "When a suspicious event is detected, the AI service prepares an incident payload that includes camera ID, district, event type, confidence, violence score when available, weapon flag, alert level, snapshot path, clip path, and metadata. This payload is sent to the Laravel API endpoint for storage."
    )
    doc.add_heading("1.5.4 Backend and Database Phase", level=3)
    doc.add_paragraph(
        "The Laravel backend receives incident requests, validates the data, creates or updates camera records, stores incident records in MySQL, and exposes web pages for dashboard, stream, cameras, incidents, evidence, and reports."
    )
    doc.add_heading("1.5.5 Operator Dashboard Phase", level=3)
    doc.add_paragraph(
        "The React dashboard presents recent incidents, alert levels, registered cameras, evidence links, and trend information. It allows an operator to inspect incidents quickly and review snapshots or video clips."
    )

    doc.add_heading("1.6 Documentation Organization", level=2)
    doc.add_paragraph("The rest of this documentation is organized as follows:")
    add_bullets(
        doc,
        [
            "Chapter 2 presents the background and related work for computer vision, violence detection, weapon detection, and security monitoring systems.",
            "Chapter 3 explains the proposed system architecture and design decisions.",
            "Chapter 4 describes implementation details for the backend, frontend dashboard, AI services, API integration, and evidence storage.",
            "Chapter 5 discusses testing methodology, test cases, and expected results.",
            "Chapter 6 presents the conclusion and possible future work.",
            "Appendix A provides a user manual for running and using the system.",
            "Appendix B lists references used to support the documentation.",
        ],
    )
    doc.add_page_break()


def other_chapters(doc):
    doc.add_heading("Chapter 2: Background and Related Work", level=1)
    doc.add_heading("2.1 Computer Vision in Public Safety", level=2)
    doc.add_paragraph(
        "Computer vision is widely used to extract information from visual data. In public-safety applications, it can support object detection, activity recognition, crowd analysis, and abnormal-event detection. These techniques are useful when camera networks produce more visual information than operators can comfortably monitor."
    )
    doc.add_heading("2.2 Violence Detection", level=2)
    doc.add_paragraph(
        "Violence detection aims to identify physical actions that may indicate fighting, assault, or aggressive movement. Video-based methods usually analyze sequences of frames because motion over time is important for distinguishing normal actions from violent ones."
    )
    doc.add_heading("2.3 Weapon Detection", level=2)
    doc.add_paragraph(
        "Weapon detection focuses on identifying dangerous objects such as guns or knives in video frames. Object-detection models such as YOLO are suitable for this task because they can localize objects and return confidence scores quickly enough for near-real-time applications."
    )
    doc.add_heading("2.4 Existing Monitoring Limitations", level=2)
    add_table(
        doc,
        ["Monitoring Approach", "Advantages", "Limitations"],
        [
            ["Manual CCTV observation", "Simple to deploy; human judgment is available", "Hard to scale; events can be missed; operators become overloaded"],
            ["Recording-only cameras", "Provides evidence after an event", "Does not support early intervention"],
            ["AI-assisted monitoring", "Continuous analysis; faster alerts; structured evidence", "Requires model training, threshold tuning, and responsible deployment"],
        ],
        widths=[2500, 3300, 3400],
    )
    add_caption(doc, "Table 2.1: Comparison of monitoring approaches")
    doc.add_page_break()

    doc.add_heading("Chapter 3: System Architecture and Design", level=1)
    doc.add_heading("3.1 Architecture Overview", level=2)
    doc.add_paragraph(
        "The system is designed as separated services. Laravel handles authentication, web routes, API routes, database models, and incident storage. React/Vite powers interactive dashboard screens. Python services handle the AI tasks and expose MJPEG streams and control endpoints."
    )
    doc.add_heading("3.2 Main System Modules", level=2)
    add_table(
        doc,
        ["Module", "Role in the System"],
        [
            ["Camera Module", "Represents monitored sources and keeps camera identifiers, district data, status, and stream information."],
            ["AI Services", "Analyze frames, detect violence or weapons, save evidence, and post incidents to the backend."],
            ["Incident API", "Receives validated incident data from AI services and stores it in the database."],
            ["Dashboard", "Displays recent incidents, alert mix, trends, cameras, and evidence for operators."],
            ["Evidence Storage", "Stores snapshots and video clips under the public incident storage path."],
        ],
        widths=[2600, 6600],
    )
    add_caption(doc, "Table 3.1: Main system modules")
    doc.add_heading("3.3 Data Flow", level=2)
    doc.add_paragraph(
        "The data flow begins with a camera frame. The AI service reads the frame, applies its model, annotates the stream when needed, and checks confidence thresholds. If an incident is detected, the service saves a snapshot and clip, builds a JSON payload, and sends it to the backend. The backend stores the record and makes it visible in dashboard pages."
    )
    doc.add_page_break()

    doc.add_heading("Chapter 4: Implementation", level=1)
    doc.add_heading("4.1 Backend Implementation", level=2)
    doc.add_paragraph(
        "The backend is implemented with Laravel. It provides web routes for dashboard pages and API routes for receiving incidents. The main models are Camera, Incident, and User. The IncidentController validates incoming data and stores the event with its alert level and evidence paths."
    )
    doc.add_heading("4.2 Frontend Dashboard Implementation", level=2)
    doc.add_paragraph(
        "The frontend uses React with Vite inside the Laravel project. It provides dashboard components for incident tables, camera lists, stream controls, alert badges, and evidence links. The dashboard is designed for operators who need fast scanning and direct access to incident details."
    )
    doc.add_heading("4.3 AI Service Implementation", level=2)
    doc.add_paragraph(
        "The Python AI services use OpenCV for frame capture and streaming. The weapon service loads a YOLO model and checks object-detection confidence. The violence service loads a Keras model and predicts whether recent frames indicate violent activity. Both services save snapshots and clips under the public storage incident folders."
    )
    doc.add_heading("4.4 API Integration", level=2)
    add_table(
        doc,
        ["Endpoint", "Method", "Purpose"],
        [
            ["/api/incidents", "GET", "Return stored incidents with related camera data."],
            ["/api/incidents", "POST", "Receive new incident records from AI services."],
            ["AI service /status", "GET", "Return camera/model status for a service."],
            ["AI service /stream", "GET", "Return MJPEG stream for monitoring."],
            ["AI service /start and /stop", "POST", "Control tracking mode in the AI service."],
        ],
        widths=[3300, 1300, 4600],
    )
    add_caption(doc, "Table 4.1: Important system endpoints")
    doc.add_page_break()

    doc.add_heading("Chapter 5: Testing", level=1)
    doc.add_heading("5.1 Testing Methodology", level=2)
    doc.add_paragraph(
        "Testing focuses on validating the complete workflow from camera stream to dashboard display. The system should be tested module by module and then as an integrated pipeline."
    )
    doc.add_heading("5.2 Test Cases", level=2)
    add_table(
        doc,
        ["Test Case", "Expected Result"],
        [
            ["Open the Laravel dashboard", "The login/dashboard pages load without errors."],
            ["Start an AI service", "The service returns a valid status response and stream endpoint."],
            ["Submit an incident to /api/incidents", "The backend validates and stores the incident record."],
            ["Detect a weapon event", "A critical alert is created with weapon_detected set to true."],
            ["Detect a violence event", "A high alert is created with a violence score and saved evidence."],
            ["Open incident details", "The operator can view event data, camera data, snapshot, clip, and metadata."],
        ],
        widths=[3900, 5300],
    )
    add_caption(doc, "Table 5.1: Functional test cases")
    doc.add_heading("5.3 Testing Results", level=2)
    doc.add_paragraph(
        "The expected testing result is that incidents are created correctly, evidence files are reachable through public storage, alert levels are displayed with the correct priority, and the dashboard refreshes recent incidents for operator review."
    )
    doc.add_page_break()

    doc.add_heading("Chapter 6: Conclusion and Future Work", level=1)
    doc.add_heading("6.1 Conclusion", level=2)
    doc.add_paragraph(
        "This project presents an AI-assisted violence and weapon detection system that connects computer vision models with a practical web monitoring platform. It supports faster incident awareness, organized evidence storage, and clearer operator review. The system is especially relevant for improving public-safety monitoring in Egypt, while also being adaptable to global environments that require early detection of violent or criminal actions."
    )
    doc.add_heading("6.2 Future Work", level=2)
    add_bullets(
        doc,
        [
            "Combine violence and weapon detection into one camera pipeline to avoid opening the same camera from multiple processes.",
            "Improve the models using larger and more diverse datasets from different lighting conditions, camera angles, and environments.",
            "Add role-based access control for different operator and administrator permissions.",
            "Add real-time notification channels such as SMS, email, or mobile push notifications.",
            "Deploy the system on edge devices or local servers near camera networks to reduce latency.",
            "Add map-based incident visualization for districts and public areas.",
        ],
    )
    doc.add_page_break()

    doc.add_heading("Appendix A: User Manual", level=1)
    doc.add_heading("A.1 Running the Web Application", level=2)
    add_numbered(
        doc,
        [
            "Install PHP and Node dependencies using composer install and npm install.",
            "Create the environment file and configure the MySQL database.",
            "Run migrations and seeders, then create the public storage link.",
            "Start the Laravel server and Vite development server.",
            "Open the login page and access the operator dashboard.",
        ],
    )
    doc.add_heading("A.2 Running the AI Services", level=2)
    add_numbered(
        doc,
        [
            "Install Python requirements for the AI service environment.",
            "Place the weapon and violence model files in the configured storage paths.",
            "Run the weapon detection service or the violence detection service.",
            "Check the /status endpoint to confirm that the camera is open.",
            "Use the dashboard stream page and incident pages to review detections.",
        ],
    )
    doc.add_page_break()

    doc.add_heading("Appendix B: References", level=1)
    refs = [
        "Laravel Documentation, \"Laravel Web Framework Documentation\", Laravel, 2026.",
        "React Documentation, \"React User Interface Library Documentation\", Meta, 2026.",
        "OpenCV Documentation, \"Open Source Computer Vision Library\", OpenCV, 2026.",
        "Ultralytics, \"YOLO Object Detection Documentation\", Ultralytics, 2026.",
        "TensorFlow/Keras Documentation, \"Keras Deep Learning API\", Google, 2026.",
    ]
    add_numbered(doc, refs)


def main():
    draw_architecture()
    doc = Document()
    style_document(doc)
    add_header_footer(doc.sections[0])
    cover(doc)
    front_matter(doc)
    chapter1(doc)
    other_chapters(doc)
    for section in doc.sections:
        add_header_footer(section)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
