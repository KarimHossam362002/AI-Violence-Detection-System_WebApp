from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


BASE_DIR = Path(__file__).resolve().parents[1]
OUT_DIR = BASE_DIR / "generated_docx_assets"
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = BASE_DIR / "AI Violence Detection - Chapters 4 5 6.docx"


def font(size=28, bold=False):
    candidates = [
        "C:/Windows/Fonts/timesbd.ttf" if bold else "C:/Windows/Fonts/times.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrap_text(draw, text, font_obj, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        test = f"{current} {word}".strip()
        if draw.textbbox((0, 0), test, font=font_obj)[2] <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def center_text(draw, box, text, font_obj, fill=(20, 20, 20), max_lines=3):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font_obj, x2 - x1 - 28)[:max_lines]
    heights = [draw.textbbox((0, 0), line, font=font_obj)[3] for line in lines]
    total = sum(heights) + (len(lines) - 1) * 8
    y = y1 + ((y2 - y1) - total) / 2
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=font_obj)[2]
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=font_obj, fill=fill)
        y += h + 8


def arrow(draw, start, end, fill=(58, 82, 110), width=4):
    draw.line([start, end], fill=fill, width=width)
    ex, ey = end
    sx, sy = start
    if ex >= sx:
        head = [(ex, ey), (ex - 16, ey - 9), (ex - 16, ey + 9)]
    else:
        head = [(ex, ey), (ex + 16, ey - 9), (ex + 16, ey + 9)]
    draw.polygon(head, fill=fill)


def save_architecture_figure():
    img = Image.new("RGB", (1400, 700), "white")
    d = ImageDraw.Draw(img)
    title_font = font(34, True)
    label_font = font(24, True)
    small_font = font(20)
    d.text((40, 30), "AI Violence Detection System Architecture", font=title_font, fill=(15, 40, 70))
    boxes = [
        ((70, 170, 270, 310), "CCTV / Webcam", "Video frames"),
        ((355, 115, 590, 255), "Weapon Service", "YOLO model, port 5000"),
        ((355, 305, 590, 445), "Violence Service", "Keras model, port 5001"),
        ((675, 210, 900, 350), "Laravel API", "POST /api/incidents"),
        ((985, 115, 1190, 255), "MySQL", "Cameras and incidents"),
        ((985, 305, 1190, 445), "React Dashboard", "Live stream and alerts"),
    ]
    fills = [(230, 242, 255), (235, 248, 239), (255, 243, 230), (239, 239, 255), (245, 245, 245), (234, 246, 250)]
    for (box, title, subtitle), fill in zip(boxes, fills):
        d.rounded_rectangle(box, radius=24, fill=fill, outline=(75, 95, 120), width=3)
        center_text(d, (box[0], box[1] + 10, box[2], box[3] - 32), title, label_font)
        center_text(d, (box[0] + 8, box[3] - 48, box[2] - 8, box[3] - 8), subtitle, small_font, fill=(70, 70, 70), max_lines=2)
    arrow(d, (270, 240), (355, 185))
    arrow(d, (270, 240), (355, 375))
    arrow(d, (590, 185), (675, 255))
    arrow(d, (590, 375), (675, 305))
    arrow(d, (900, 255), (985, 185))
    arrow(d, (900, 305), (985, 375))
    arrow(d, (985, 375), (900, 330))
    path = OUT_DIR / "figure_4_1_architecture.png"
    img.save(path)
    return path


def save_dataset_figure():
    img = Image.new("RGB", (1200, 680), "white")
    d = ImageDraw.Draw(img)
    title = font(32, True)
    label = font(22, True)
    small = font(19)
    d.text((40, 30), "Dataset Sources Used in the Project", font=title, fill=(15, 40, 70))
    categories = [
        ("Violence / fight videos", 5, (80, 145, 210)),
        ("Anomaly / CCTV videos", 2, (225, 125, 75)),
        ("Weapon images", 4, (95, 170, 120)),
    ]
    total = sum(item[1] for item in categories)
    x0, y0, bar_w, bar_h = 90, 165, 870, 70
    current = x0
    for name, value, color in categories:
        width = int(bar_w * value / total)
        d.rectangle((current, y0, current + width, y0 + bar_h), fill=color)
        center_text(d, (current, y0, current + width, y0 + bar_h), str(value), label, fill=(255, 255, 255), max_lines=1)
        current += width
    d.rectangle((x0, y0, x0 + bar_w, y0 + bar_h), outline=(40, 40, 40), width=2)
    y = 290
    for name, value, color in categories:
        d.rectangle((90, y + 4, 120, y + 34), fill=color)
        d.text((135, y), f"{name}: {value} sources", font=small, fill=(30, 30, 30))
        y += 55
    notes = [
        "Violence datasets support action recognition and fight/non-fight classification.",
        "Weapon datasets support firearm, knife, stick, and rod detection.",
        "The final data pool combines public academic datasets, Kaggle datasets, and Roboflow/Universe datasets.",
    ]
    y = 470
    for note in notes:
        d.text((90, y), "- " + note, font=small, fill=(40, 40, 40))
        y += 42
    path = OUT_DIR / "figure_4_2_dataset_sources.png"
    img.save(path)
    return path


def save_preprocessing_figure():
    img = Image.new("RGB", (1300, 520), "white")
    d = ImageDraw.Draw(img)
    title = font(32, True)
    label = font(23, True)
    small = font(18)
    d.text((40, 25), "Preprocessing Pipeline", font=title, fill=(15, 40, 70))
    steps = [
        ("Read frames", "OpenCV captures camera/video frames"),
        ("Resize", "640x480 stream, 224x224 model input"),
        ("Normalize", "MobileNetV2 preprocess_input"),
        ("Window frames", "20 frames for violence sequence"),
        ("Predict", "Keras violence or YOLO weapon model"),
    ]
    x = 55
    y = 155
    for i, (a, b) in enumerate(steps):
        box = (x, y, x + 200, y + 135)
        d.rounded_rectangle(box, radius=18, fill=(240, 246, 252), outline=(65, 90, 120), width=3)
        center_text(d, (box[0], box[1] + 10, box[2], box[1] + 64), a, label, max_lines=2)
        center_text(d, (box[0] + 8, box[1] + 68, box[2] - 8, box[3] - 8), b, small, fill=(70, 70, 70), max_lines=3)
        if i < len(steps) - 1:
            arrow(d, (x + 200, y + 68), (x + 255, y + 68), width=4)
        x += 255
    path = OUT_DIR / "figure_4_3_preprocessing.png"
    img.save(path)
    return path


def save_results_figure():
    img = Image.new("RGB", (1250, 680), "white")
    d = ImageDraw.Draw(img)
    title = font(32, True)
    label = font(22, True)
    small = font(19)
    d.text((40, 25), "Incident Output Flow", font=title, fill=(15, 40, 70))
    columns = [
        ("Detection", ["Violence score", "Weapon flag", "Confidence"]),
        ("Evidence", ["Snapshot .jpg", "Clip .mp4", "H.264 conversion"]),
        ("API payload", ["camera_id", "event_type", "alert_level"]),
        ("Dashboard", ["Critical/high cards", "Incident table", "Evidence links"]),
    ]
    x = 60
    for title_text, items in columns:
        d.rounded_rectangle((x, 120, x + 250, 500), radius=18, fill=(248, 248, 248), outline=(90, 90, 90), width=2)
        d.rectangle((x, 120, x + 250, 175), fill=(224, 235, 246))
        center_text(d, (x, 123, x + 250, 170), title_text, label)
        y = 220
        for item in items:
            d.ellipse((x + 24, y + 5, x + 38, y + 19), fill=(50, 120, 185))
            d.text((x + 55, y), item, font=small, fill=(40, 40, 40))
            y += 72
        if x < 900:
            arrow(d, (x + 250, 310), (x + 295, 310), width=4)
        x += 300
    path = OUT_DIR / "figure_4_4_incident_flow.png"
    img.save(path)
    return path


def save_run_workflow_figure():
    img = Image.new("RGB", (1300, 620), "white")
    d = ImageDraw.Draw(img)
    title = font(32, True)
    label = font(23, True)
    small = font(18)
    d.text((40, 25), "Application Running Workflow", font=title, fill=(15, 40, 70))
    steps = [
        ("1", "Start Laravel", "php artisan serve"),
        ("2", "Start Vite", "npm run dev"),
        ("3", "Start AI service", "python ai_service/..."),
        ("4", "Open dashboard", "127.0.0.1:8000/login"),
        ("5", "Load stream", "Weapon or violence port"),
        ("6", "Start AI", "Detect, save, alert"),
    ]
    x = 55
    y = 155
    for n, title_text, subtitle in steps:
        box = (x, y, x + 180, y + 150)
        d.rounded_rectangle(box, radius=22, fill=(239, 248, 244), outline=(65, 105, 90), width=3)
        d.ellipse((x + 18, y + 18, x + 58, y + 58), fill=(65, 130, 95))
        center_text(d, (x + 18, y + 15, x + 58, y + 58), n, label, fill=(255, 255, 255), max_lines=1)
        center_text(d, (x + 10, y + 56, x + 170, y + 100), title_text, label, max_lines=2)
        center_text(d, (x + 8, y + 105, x + 172, y + 145), subtitle, small, fill=(70, 70, 70), max_lines=2)
        if n != "6":
            arrow(d, (x + 180, y + 75), (x + 225, y + 75), width=4)
        x += 210
    path = OUT_DIR / "figure_5_1_running_workflow.png"
    img.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def set_document_style(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(8)
    for name, size in [("Heading 1", 18), ("Heading 2", 16), ("Heading 3", 14)]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
    caption = styles.add_style("CaptionSimple", 1)
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(12)
    caption.font.italic = False


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    run.append(t)
    fld.append(run)
    p._p.append(fld)


def para(doc, text="", bold_start=None):
    p = doc.add_paragraph()
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def caption(doc, text):
    p = doc.add_paragraph(style="CaptionSimple")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = False


def image(doc, path, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))


def chapter_title_page(doc, number_word, title):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(45)
    r = p.add_run(f"Chapter {number_word}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(32)
    r.bold = True
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(45)
    r2 = p2.add_run(title)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(24)
    r2.bold = True
    doc.add_page_break()


def build_doc():
    figs = {
        "arch": save_architecture_figure(),
        "datasets": save_dataset_figure(),
        "pre": save_preprocessing_figure(),
        "results": save_results_figure(),
        "run": save_run_workflow_figure(),
    }

    doc = Document()
    set_document_style(doc)
    add_page_number(doc.sections[0])

    chapter_title_page(doc, "Four", "System Implementation and Results")
    doc.add_heading("4.1 Dataset", level=1)
    para(doc, "The proposed AI Violence Detection System depends on two related dataset groups. The first group contains CCTV and real-world video clips for violence, fighting, and abnormal behavior recognition. The second group contains annotated weapon images for firearms, knives, sticks, and rods. Combining both groups helps the system detect the action itself and also recognize dangerous objects that may increase the alert level.")
    image(doc, figs["datasets"], 6.1)
    caption(doc, "Figure 4.1: Dataset source groups used in the project")
    para(doc, "The datasets used in this project are listed in Table 4.1. The video datasets were used to support the violence detection model, while the object detection datasets were used to train and validate weapon detection behavior.")
    add_table(
        doc,
        ["No.", "Dataset / Source", "Main Use"],
        [
            ["1", "Mohamed et al. (2019), Violence Recognition from Videos using Deep Learning Techniques", "Violence and non-violence video recognition"],
            ["2", "Atharva et al. (2025), Real-Time CCTV Violence Detection Automation System", "Real-time CCTV violence examples"],
            ["3", "Aremu et al. (2024), Smart-City CCTV Violence Detection Dataset (SCVD)", "Smart-city surveillance violence detection"],
            ["4", "Fransco Perez Hernandez et al. (2020), OD-WeaponDetection: Knife Classification", "Knife object detection/classification"],
            ["5", "Ali Alrabeei et al. (2024), CCTV Guns Online", "Gun detection in CCTV scenes"],
            ["6", "Violence Detection et al. (2024), Stick/Rod Detection", "Stick and rod object detection"],
            ["7", "S Kumar et al. (2025), Firearms Detection", "Firearm detection"],
            ["8", "Dania Zehra et al. (2025), RWF-2000", "Fight/non-fight video classification"],
            ["9", "Shrey Joshi et al. (2021), NTU CCTV-Fights Dataset", "CCTV fight detection"],
            ["10", "Avdhesh Chaudhary et al. (2023), UCF Crime Dataset", "CCTV anomaly and violent crime scenes"],
        ],
        widths=[0.45, 4.15, 1.65],
    )
    caption(doc, "Table 4.1: Datasets used in the project")

    doc.add_heading("4.2 Description of Software Tools", level=1)
    para(doc, "Several tools were used to implement the web application, the artificial intelligence services, and the live monitoring dashboard. The most important AI and computer vision tools are shown in Table 4.2.")
    add_table(
        doc,
        ["Tool", "Version", "Use in the Project"],
        [
            ["Flask", "3.1.3", "Builds the lightweight Python APIs for weapon and violence detection streams."],
            ["TensorFlow", "2.21.0", "Loads and runs the trained violence detection model."],
            ["Keras", "3.14.0", "Provides the model interface used for the violence classification network."],
            ["OpenCV", "4.13.0.92", "Captures frames, resizes images, draws overlays, saves snapshots, and records clips."],
            ["Laravel", "12.x", "Provides authentication, database models, dashboard pages, and incident API endpoints."],
            ["React / Vite", "Project frontend stack", "Builds the operator dashboard, stream controls, and incident table."],
            ["MySQL", "Project database", "Stores cameras, incidents, confidence scores, alert levels, and evidence links."],
        ],
        widths=[1.25, 1.1, 4.0],
    )
    caption(doc, "Table 4.2: Main software tools used in implementation")

    doc.add_heading("4.3 Experimental Setup and Results", level=1)
    para(doc, "The implemented system is divided into a Laravel web backend, a React dashboard, and two Python AI services. The weapon service runs on port 5000 and the violence service runs on port 5001. Each service provides a live MJPEG stream and start/stop endpoints for controlling analysis from the dashboard.")
    image(doc, figs["arch"], 6.2)
    caption(doc, "Figure 4.2: System implementation architecture")

    doc.add_heading("4.3.1 Data Preprocessing", level=2)
    para(doc, "For violence detection, frames are captured from the camera or video source, resized, converted to floating point values, and passed through the MobileNetV2 preprocessing function. The system collects a short sequence of 20 frames before prediction because violence is an action that depends on motion over time. For weapon detection, frames are resized and passed to the YOLO model, which returns bounding boxes and confidence scores.")
    image(doc, figs["pre"], 6.2)
    caption(doc, "Figure 4.3: Data preprocessing pipeline")
    bullet(doc, "Frame size used in the live stream: 640 x 480 pixels.")
    bullet(doc, "Violence model input size: 224 x 224 pixels.")
    bullet(doc, "Violence decision threshold: 0.60.")
    bullet(doc, "Weapon confidence threshold: 0.65.")

    doc.add_heading("4.3.2 Model Training", level=2)
    para(doc, "The violence model follows a deep learning sequence-classification approach. MobileNetV2 is used as a feature extractor for individual frames, and the extracted features are passed to a trained Keras model saved as cctvmodel_advanced.keras. This design reduces the cost of learning visual features from scratch and allows the classifier to focus on the final violence/non-violence decision.")
    para(doc, "The weapon model uses a YOLO object detection model saved as weapon_detection_yolov11m.pt. YOLO is suitable for real-time detection because it predicts object locations and classes in one pass. In the implementation, the service filters firearm-related class names, tracks detections using ByteTrack, and starts evidence recording only after consecutive hits to reduce false alerts.")

    doc.add_heading("4.3.3 Evaluation Metrics", level=2)
    para(doc, "The following metrics are suitable for evaluating the project models:")
    bullet(doc, "Accuracy = correctly classified samples / total samples.")
    bullet(doc, "Precision = true positives / (true positives + false positives).")
    bullet(doc, "Recall = true positives / (true positives + false negatives).")
    bullet(doc, "F1-score = 2 x (precision x recall) / (precision + recall).")
    bullet(doc, "Mean Average Precision (mAP) is used for weapon object detection because it measures both classification and localization quality.")
    bullet(doc, "Response time is used at system level to check whether the stream and alert creation are fast enough for live monitoring.")

    doc.add_heading("4.3.4 Results Presentation", level=2)
    para(doc, "The implemented system produces visible and stored outputs when a suspicious event is detected. The AI service draws detection overlays on the stream, saves a snapshot, records a short video clip, and sends an incident payload to the Laravel API. Laravel stores the event in the database and the dashboard displays it in the incident table.")
    image(doc, figs["results"], 6.1)
    caption(doc, "Figure 4.4: Incident output flow")
    add_table(
        doc,
        ["Test Case", "Expected Result", "Observed Implementation Result"],
        [
            ["Open weapon stream", "MJPEG stream is available on port 5000", "Implemented through /stream endpoint."],
            ["Open violence stream", "MJPEG stream is available on port 5001", "Implemented through /stream endpoint."],
            ["Start/stop AI", "Dashboard can control Python analysis", "Implemented through /start and /stop endpoints."],
            ["Weapon detection", "Critical incident is created when weapon confidence passes threshold", "Payload includes weapon_detected=true and alert_level=critical."],
            ["Violence detection", "High incident is created when violence score passes threshold", "Payload includes violence_score and alert_level=high."],
            ["Evidence saving", "Snapshot and clip are saved for review", "Files are saved under storage/app/public/incidents."],
            ["Dashboard update", "New incidents appear without manual refresh", "Dashboard polls latest incidents every five seconds."],
        ],
        widths=[1.7, 2.35, 2.3],
    )
    caption(doc, "Table 4.3: Functional results of the implemented system")

    doc.add_heading("4.3.5 Comparative Analysis", level=2)
    para(doc, "Compared with a simple single-model violence classifier, the proposed system adds a second weapon detection service. This improves practical safety monitoring because a scene can be dangerous even when the motion pattern is not clearly classified as violence. Compared with an offline notebook experiment, the implemented system provides a complete application pipeline: live stream, inference service, evidence storage, incident API, database records, and dashboard review.")
    add_table(
        doc,
        ["Aspect", "Single Offline Model", "Proposed System"],
        [
            ["Input", "Prepared videos or images", "Live webcam/CCTV stream"],
            ["Detection scope", "Usually one task", "Violence detection and weapon detection"],
            ["Output", "Prediction score", "Alert level, snapshot, clip, dashboard record"],
            ["Deployment", "Experiment-only environment", "Flask AI services connected to Laravel dashboard"],
            ["Review", "Manual checking", "Incident archive with evidence links"],
        ],
        widths=[1.4, 2.4, 2.55],
    )
    caption(doc, "Table 4.4: Comparison between offline detection and the proposed implemented system")

    doc.add_heading("4.4 Discussion", level=1)
    para(doc, "The implementation shows that deep learning can be connected with a practical surveillance dashboard. The violence model handles action recognition through frame sequences, while the weapon model handles object detection. The dashboard and incident API make the system more useful than a standalone model because the operator can monitor events, inspect evidence, and review alert history.")
    para(doc, "The main limitation is that real-time performance depends on the camera source, CPU/GPU capability, and model size. Also, lighting, occlusion, crowd density, and camera angle can affect detection confidence. For this reason, the system stores confidence values and evidence files instead of relying only on a binary alert.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    chapter_title_page(doc, "Five", "Run the Application")
    doc.add_heading("5.1 Running the Web Application", level=1)
    para(doc, "This chapter explains how to run and use the application. The project can be started locally by running the Laravel backend, the Vite frontend, and one of the Python AI services.")
    image(doc, figs["run"], 6.2)
    caption(doc, "Figure 5.1: Main running workflow")
    numbered(doc, "Install PHP and Node dependencies using composer install and npm install.")
    numbered(doc, "Create the .env file, configure MySQL, generate the application key, and run database migrations.")
    numbered(doc, "Start Laravel on http://127.0.0.1:8000 and start Vite for frontend assets.")
    numbered(doc, "Run the weapon service or violence service from the ai_service folder.")
    numbered(doc, "Open the login page and enter the dashboard.")

    doc.add_heading("5.2 Running the AI Services", level=1)
    para(doc, "The weapon service is started by running ai_service/weapon_detection_service.py. It loads the YOLO model from storage/app/models/weapon/weapon_detection_yolov11m.pt and exposes a stream on http://127.0.0.1:5000/stream.")
    para(doc, "The violence service is started by running ai_service/violence_detection_service.py. It loads the Keras model from storage/app/models/violence/cctvmodel_advanced.keras and exposes a stream on http://127.0.0.1:5001/stream.")
    add_table(
        doc,
        ["Service", "Local URL", "Purpose"],
        [
            ["Laravel backend", "http://127.0.0.1:8000", "Dashboard and incident API"],
            ["Weapon AI", "http://127.0.0.1:5000", "Weapon detection stream"],
            ["Violence AI", "http://127.0.0.1:5001", "Violence detection stream"],
            ["MySQL", "127.0.0.1:3306", "Application database"],
        ],
        widths=[1.6, 2.3, 2.4],
    )
    caption(doc, "Table 5.1: Main local services")

    doc.add_heading("5.3 Dashboard Usage", level=1)
    para(doc, "After logging in, the operator can choose either weapon detection or violence detection, load the stream, and press Start AI. The dashboard shows the active stream, current system status, number of cameras, and recent incidents. When the AI service detects an event, it posts the incident to Laravel and the dashboard displays the alert with confidence and evidence links.")
    bullet(doc, "Start webcam: opens the browser camera preview.")
    bullet(doc, "Load stream: loads the selected Python AI stream into the dashboard.")
    bullet(doc, "Start AI: sends a request to the selected Python service to begin analysis.")
    bullet(doc, "Stop AI: pauses analysis and closes the current recording if needed.")
    bullet(doc, "Incident archive: lists event type, camera, confidence, detected time, snapshot, and clip.")

    doc.add_heading("5.4 Evidence Review", level=1)
    para(doc, "Each incident can contain a snapshot and a recorded video clip. These files are saved in the public storage path and linked from the dashboard. This makes the alert review process easier because the operator can inspect what caused the detection instead of depending only on the model score.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    chapter_title_page(doc, "Six", "Conclusion and Future Work")
    doc.add_heading("6.1 Conclusion", level=1)
    para(doc, "This project implemented an AI-based violence detection and weapon detection system for CCTV-style monitoring. The final system combines Python AI services, OpenCV video processing, TensorFlow/Keras violence classification, YOLO weapon detection, a Laravel backend, a MySQL database, and a React dashboard.")
    para(doc, "The implementation achieved the main project goal: detecting suspicious events from a live stream, creating incident records, saving evidence, and presenting alerts to the user through a web dashboard. The separation between the Laravel backend and Python AI services also makes the system easier to maintain because model inference and web application logic are handled independently.")

    doc.add_heading("6.2 Limitations", level=1)
    bullet(doc, "The system performance depends on camera quality, lighting conditions, and available hardware.")
    bullet(doc, "Running two OpenCV services on the same webcam may fail on Windows because one camera is usually locked by one process.")
    bullet(doc, "The current dashboard uses polling every five seconds; real-time WebSocket alerts would be faster.")
    bullet(doc, "Final accuracy values should be reported after running the trained models on a fixed test split and saving the evaluation logs.")

    doc.add_heading("6.3 Future Work", level=1)
    bullet(doc, "Create a combined AI service that opens the camera once and runs both violence and weapon models on the same frames.")
    bullet(doc, "Add WebSocket notifications for instant dashboard alerts.")
    bullet(doc, "Improve the dataset by adding more local CCTV scenes with different lighting, camera angles, and crowd densities.")
    bullet(doc, "Add role-based incident review actions such as confirmed, false alarm, ignored, and escalated.")
    bullet(doc, "Deploy the system on an edge device or GPU-enabled server for better real-time performance.")
    bullet(doc, "Add model evaluation reports directly to the dashboard, including confusion matrix, precision, recall, F1-score, and mAP.")

    doc.add_heading("6.4 References", level=1)
    refs = [
        "Mohamed et al., \"Violence Recognition from Videos using Deep Learning Techniques,\" 2019. https://ieeexplore.ieee.org/abstract/document/9014714",
        "Atharva et al., \"Real-Time CCTV Violence Detection Automation System,\" 2025. https://drive.google.com/file/d/1dXrt-dXBmjBao3yIBLhqTv4e3iKti7Hw/view",
        "Aremu et al., \"Smart-City CCTV Violence Detection Dataset (SCVD),\" 2024. https://www.kaggle.com/datasets/toluwaniaremu/smartcity-cctv-violence-detection-dataset-scvd",
        "Fransco Perez Hernandez et al., \"OD-WeaponDetection: Knife Classification,\" 2020. https://datasetninja.com/od-weapon-detection-knife-classification",
        "Ali Alrabeei et al., \"CCTV Guns Online,\" 2024. https://universe.roboflow.com/ali-alrabeei-qb1zy/cctv-guns-online",
        "Violence Detection et al., \"Stick/Rod Detection,\" 2024. https://universe.roboflow.com/violence-detection-rvh0k/stick-rod-detection",
        "S Kumar et al., \"Firearms Detection,\" 2025. https://universe.roboflow.com/s-kumar-baywt/firearms-detection-5wjh7",
        "Dania Zehra et al., \"RWF-2000,\" 2025. https://www.kaggle.com/datasets/daniazehra/rwf-2000",
        "Shrey Joshi et al., \"NTU CCTV-Fights Dataset,\" 2021. https://www.kaggle.com/datasets/shreyj1729/cctv-fights-dataset/data",
        "Avdhesh Chaudhary et al., \"UCF Crime Dataset,\" 2023. https://www.kaggle.com/datasets/webadvisor/real-time-anomaly-detection-in-cctv-surveillance",
    ]
    for ref in refs:
        bullet(doc, ref)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_doc())
